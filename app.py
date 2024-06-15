import gradio as gr
from HybridRetriever import HybridRetriever
from ChatEngine import ChatEngine
from configs import *

import warnings
warnings.simplefilter("ignore", UserWarning)
warnings.simplefilter("ignore", FutureWarning)
from llama_index.core.readers import SimpleDirectoryReader
from llama_index.retrievers.bm25 import BM25Retriever 
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core import VectorStoreIndex, Document
from llama_index.llms.ollama import Ollama
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings
from llama_index.core.node_parser import SentenceSplitter
import fitz
from docx import Document as DocxDocument
import csv

llm = Ollama(model=MODEL_NAME, request_timeout=1200.0, context_window=CONTEXT_WINDOW)
embedding = OllamaEmbedding(model_name="nomic-embed-text:latest")

Settings.llm = llm
Settings.embed_model = embedding

def process_file(file):
    file_extension = file.name.split(".")[-1].lower()

    if file_extension == 'txt':
        with open(file.name, 'r', encoding='utf-8') as f:
            text = f.read()

    elif file_extension == 'csv':
        with open(file.name, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = '\n'.join(','.join(row) for row in reader)

    elif file_extension == 'pdf':
        pdf_document = fitz.open(file.name, filetype=file_extension)
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        pdf_document.close()

    elif file_extension == 'docx':
        docx_document = DocxDocument(file.name)
        text = ""
        for paragraph in docx_document.paragraphs:
            text += paragraph.text + "\n"

    return [Document(text=text)]


def process_and_respond(file, question):
    global llm

    documents = process_file(file)
    # eger huggignface space dışında kullanılıyorsa, 
    # alttaki kodları kullan, huggingface ise üsttekini
    # documents = SimpleDirectoryReader("./data").load_data()


    text_splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    vector_index = VectorStoreIndex.from_documents(
        documents, transformations=[text_splitter], embed_model=Settings.embed_model, show_progress=True
    )

    bm25_retriever = BM25Retriever(nodes=documents, similarity_top_k=TOP_K, tokenizer=text_splitter.split_text)
    vector_retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=TOP_K)
    hybrid_retriever = HybridRetriever(bm25_retriever=bm25_retriever, vector_retriever=vector_retriever)

    chat_engine = ChatEngine(hybrid_retriever)

    response = chat_engine.ask_question(question, llm)

    return response


def main():
    with gr.Blocks() as demo:
        gr.Markdown("## Gradio App with Text Input, File Uploader, and Response Textbox")

        text_input = gr.Textbox(label="Enter your question here:")
        file_uploader = gr.File(label="Upload a file:")

        response_box = gr.TextArea(max_lines=50)

        submit_button = gr.Button("Submit")

        submit_button.click(
            fn=process_and_respond,
            inputs=[file_uploader, text_input],
            outputs=response_box
        )

    demo.launch()

if __name__ == "__main__":
    main()
